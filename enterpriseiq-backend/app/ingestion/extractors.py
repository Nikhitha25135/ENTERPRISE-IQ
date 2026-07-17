"""
Document extraction layer.

Each extractor takes a file path and returns a normalized structure:

    {
        "kind": "text" | "table",
        "pages": [ { "page": int, "text": str }, ... ]          # for "text"
        "tables": [ { "sheet": str, "records": list[dict] }, .. ] # for "table"
    }

Keeping "text" and "table" separate (rather than flattening everything into
one big string) is what lets the chunking layer use structure-aware
strategies, and lets a future Analytics Agent query the real DataFrame
instead of guessing numbers from embedded text fragments.
"""
import logging
import os

from app.core.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


def extract_pdf(path: str) -> dict:
    import fitz  # PyMuPDF

    pages = []
    doc = fitz.open(path)
    try:
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            if not text and settings.ocr_enabled:
                # Scanned page with no extractable text layer -> OCR it.
                text = _ocr_pdf_page(page)
            pages.append({"page": i + 1, "text": text})
    finally:
        doc.close()
    return {"kind": "text", "pages": pages}


def _ocr_pdf_page(page) -> str:
    try:
        import pytesseract
        from PIL import Image
        import io

        if settings.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img).strip()
    except Exception:
        logger.exception("OCR failed for a scanned PDF page; continuing with empty text")
        return ""


def extract_docx(path: str) -> dict:
    import docx

    document = docx.Document(path)
    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            # Preserve heading structure so chunking can split on sections.
            style = (para.style.name or "").lower()
            if style.startswith("heading"):
                parts.append(f"\n## {para.text.strip()}\n")
            else:
                parts.append(para.text.strip())

    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            parts.append(_rows_to_markdown_table(rows))

    text = "\n".join(parts)
    return {"kind": "text", "pages": [{"page": 1, "text": text}]}


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header, *body = rows
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def extract_image(path: str) -> dict:
    if not settings.ocr_enabled:
        return {"kind": "text", "pages": [{"page": 1, "text": ""}]}

    import pytesseract
    from PIL import Image

    if settings.tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

    img = Image.open(path)
    text = pytesseract.image_to_string(img).strip()
    return {"kind": "text", "pages": [{"page": 1, "text": text}]}


def extract_tabular(path: str, ext: str) -> dict:
    """Excel/CSV are kept as real structured tables, not flattened into prose.

    Each sheet's rows are stored as JSON records (for a future Analytics
    Agent to query directly with pandas), alongside a markdown rendering
    used for embedding/semantic search.
    """
    import pandas as pd

    tables = []
    if ext in (".xlsx", ".xls"):
        sheets = pd.read_excel(path, sheet_name=None)
    else:  # .csv
        sheets = {"Sheet1": pd.read_csv(path)}

    for sheet_name, df in sheets.items():
        df = df.dropna(how="all")
        records = df.fillna("").astype(str).to_dict(orient="records")
        tables.append({"sheet": sheet_name, "records": records, "columns": list(df.columns)})

    return {"kind": "table", "tables": tables}


EXTENSION_DISPATCH = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_docx,
    ".xlsx": extract_tabular,
    ".xls": extract_tabular,
    ".csv": extract_tabular,
    ".png": extract_image,
    ".jpg": extract_image,
    ".jpeg": extract_image,
    ".txt": lambda path: {
        "kind": "text",
        "pages": [{"page": 1, "text": open(path, "r", encoding="utf-8", errors="ignore").read()}],
    },
}


def extract_document(path: str, ext: str) -> dict:
    ext = ext.lower()
    extractor = EXTENSION_DISPATCH.get(ext)
    if not extractor:
        raise ValueError(f"No extractor registered for extension '{ext}'")

    if ext in (".xlsx", ".xls", ".csv"):
        return extractor(path, ext)
    return extractor(path)
